from flask import Flask, request, jsonify, send_file, render_template, after_this_request
import os
import threading
import time
import datetime
import json
import sqlite3
import shutil
import socket
import smtplib
import importlib.util
import sys
import subprocess
import platform
from pathlib import Path
import zipfile
import tempfile

from docx import Document
import pandas as pd
from PyPDF2 import PdfReader, PdfWriter

from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

from collections import defaultdict
from werkzeug.utils import secure_filename


# ============================================================
# APP
# ============================================================

app = Flask(__name__)

# لا نرفع الحد الأدنى للنشر
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024


# ============================================================
# CONFIGURATION
# ============================================================

ACCESS_CODE = os.environ.get(
    "ACCESS_CODE",
    "user2024"
)

ADMIN_CODE = os.environ.get(
    "ADMIN_CODE",
    "admin2024"
)

GMAIL_USER = os.environ.get(
    "GMAIL_USER",
    "wapoyassin08@gmail.com"
)

GMAIL_PASS = os.environ.get(
    "GMAIL_PASS",
    "ioat otqj kyte vduq"
)


# ============================================================
# ABSOLUTE PATHS
# ============================================================

BASE_DIR = Path(__file__).resolve().parent

DATA_DIR = BASE_DIR / "data"
EXCELS_DIR = DATA_DIR / "excels"
APPLICATIONS_DIR = DATA_DIR / "applications"

LOGS_DB = DATA_DIR / "logs.db"


DATA_DIR.mkdir(
    parents=True,
    exist_ok=True
)

EXCELS_DIR.mkdir(
    parents=True,
    exist_ok=True
)

APPLICATIONS_DIR.mkdir(
    parents=True,
    exist_ok=True
)


# ============================================================
# DATABASE
# ============================================================

def init_db():

    LOGS_DB.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    conn = sqlite3.connect(
        str(LOGS_DB)
    )

    cursor = conn.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS bewerber_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_name TEXT,
            event_type TEXT,
            company_num INTEGER,
            email TEXT,
            firma TEXT,
            status TEXT,
            error_msg TEXT,
            files_sent TEXT,
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS excel_uploads (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_name TEXT,
            filename TEXT,
            storage_path TEXT,
            uploaded_at TEXT DEFAULT (datetime('now'))
        )
    """)

    conn.commit()
    conn.close()


init_db()


def get_db():

    # تأكد دائمًا أن DB موجودة
    LOGS_DB.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    conn = sqlite3.connect(
        str(LOGS_DB),
        check_same_thread=False
    )

    conn.row_factory = sqlite3.Row

    return conn


# ============================================================
# GLOBAL STATE
# ============================================================

state = {

    "logged_in": False,
    "is_admin": False,

    "generating": False,
    "sending": False,

    "generated": False,

    "base_dir": None,

    "other": None,
    "extra": None,

    "scheduled_dt": None,

    "send_done": False,
    "interrupted_at": None,

    "gen_progress": 0,
    "gen_log": [],
    "gen_total": 0,

    "send_progress": 0,
    "send_log": [],
    "send_total": 0,

    "bewerbungsname": None,
    "total_companies": 0,
    

    "anschreiben_pos": 2,

    "delay": 10,
    "start": 1,

    "waiting_scheduled": False,
    "network_error": False,
    
    "direct_mode": False,
    "direct_emails": [],
    "direct_letter_path": None,
    "direct_pdf_path": None,
    
}


state_lock = threading.Lock()


# ============================================================
# DATETIME
# ============================================================

def utc_now():

    return datetime.datetime.now(
        datetime.timezone.utc
    )


# ============================================================
# LIBREOFFICE
# ============================================================

def get_libreoffice():

    system = platform.system()

    # -----------------------------
    # WINDOWS
    # -----------------------------

    if system == "Windows":

        candidates = [

            Path(
                r"C:\Program Files\LibreOffice\program\soffice.exe"
            ),

            Path(
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            ),

        ]

        for path in candidates:

            if path.exists():

                return str(path)

        found = shutil.which(
            "soffice"
        )

        if found:
            return found

        found = shutil.which(
            "libreoffice"
        )

        if found:
            return found

        raise FileNotFoundError(
            "LibreOffice غير موجود على Windows. "
            "ثبت LibreOffice أو أضف soffice.exe إلى PATH."
        )

    # -----------------------------
    # LINUX / DOCKER
    # -----------------------------

    found = shutil.which(
        "libreoffice"
    )

    if found:
        return found

    found = shutil.which(
        "soffice"
    )

    if found:
        return found

    raise FileNotFoundError(
        "LibreOffice غير موجود في السيرفر."
    )


# ============================================================
# DOCX -> PDF
# ============================================================

def convert(docx_path, pdf_path):

    docx_path = Path(
        docx_path
    ).resolve()

    pdf_path = Path(
        pdf_path
    ).resolve()

    if not docx_path.exists():

        raise FileNotFoundError(
            f"DOCX not found: {docx_path}"
        )

    pdf_path.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    soffice = get_libreoffice()

    # profile مستقل لكل عملية
    profile_dir = (
        pdf_path.parent /
        ".libreoffice_profile"
    )

    profile_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    # Windows/Linux file URL
    profile_uri = (
        profile_dir
        .resolve()
        .as_uri()
    )

    command = [

        soffice,

        "--headless",

        f"-env:UserInstallation={profile_uri}",

        "--convert-to",
        "pdf:writer_pdf_Export",

        "--outdir",
        str(pdf_path.parent),

        str(docx_path),
    ]

    try:

        result = subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=120
        )

    except subprocess.TimeoutExpired:

        raise RuntimeError(
            "LibreOffice conversion timed out."
        )

    except FileNotFoundError:

        raise RuntimeError(
            f"LibreOffice executable not found: {soffice}"
        )

    print(
        "LibreOffice stdout:",
        result.stdout
    )

    print(
        "LibreOffice stderr:",
        result.stderr
    )

    print(
        "LibreOffice return code:",
        result.returncode
    )

    generated_pdf = (
        pdf_path.parent /
        f"{docx_path.stem}.pdf"
    )

    if not generated_pdf.exists():

        raise RuntimeError(
            "LibreOffice لم ينشئ PDF.\n"
            f"Expected: {generated_pdf}\n"
            f"stdout: {result.stdout}\n"
            f"stderr: {result.stderr}"
        )

    # إذا كان الاسم المطلوب مختلفًا
    if generated_pdf != pdf_path:

        if pdf_path.exists():
            pdf_path.unlink()

        generated_pdf.replace(
            pdf_path
        )

    if not pdf_path.exists():

        raise FileNotFoundError(
            f"PDF not found after conversion: {pdf_path}"
        )

    if pdf_path.stat().st_size <= 0:

        raise IOError(
            f"PDF is empty: {pdf_path}"
        )

    # حذف profile
    try:

        shutil.rmtree(
            profile_dir,
            ignore_errors=True
        )

    except Exception:
        pass

    return pdf_path


# ============================================================
# NETWORK ERROR
# ============================================================

def is_network_error(e):

    error_text = str(e).lower()

    keywords = [

        "connection",
        "network",
        "timeout",
        "socket",
        "refused",
        "unreachable",
        "errno",
        "broken pipe",
        "reset",
        "ssl",
        "eof",
        "timed out",

    ]

    return (

        any(
            keyword in error_text
            for keyword in keywords
        )

        or

        isinstance(
            e,
            (
                socket.timeout,
                socket.gaierror,
                OSError,
            )
        )
    )


# ============================================================
# RESET
# ============================================================

def reset_state():

    with state_lock:

        state["generating"] = False
        state["sending"] = False

        state["generated"] = False

        state["base_dir"] = None

        state["other"] = None
        state["extra"] = None

        state["scheduled_dt"] = None

        state["send_done"] = False
        state["interrupted_at"] = None

        state["gen_progress"] = 0
        state["gen_log"] = []
        state["gen_total"] = 0

        state["send_progress"] = 0
        state["send_log"] = []
        state["send_total"] = 0

        state["bewerbungsname"] = None
        state["total_companies"] = 0
        
        # Direct send mode
        "direct_mode": False,
        "direct_emails": [],
        "direct_letter_path": None,
        "direct_pdf_path": None,

        state["anschreiben_pos"] = 2

        state["delay"] = 10
        state["start"] = 1

        state["waiting_scheduled"] = False
        state["network_error"] = False

        # لا نخرج المستخدم من الحساب
        state["logged_in"] = True


# ============================================================
# LOGGING
# ============================================================

def log_event(
    session_name,
    event_type,
    company_num=None,
    email=None,
    firma=None,
    status=None,
    error_msg=None,
    files_sent=None
):

    conn = None

    try:

        # تأكد من وجود الجدول
        init_db()

        conn = get_db()

        conn.execute(
            """
            INSERT INTO bewerber_logs
            (
                session_name,
                event_type,
                company_num,
                email,
                firma,
                status,
                error_msg,
                files_sent
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                session_name,
                event_type,
                company_num,
                email,
                firma,
                status,
                error_msg or "",
                json.dumps(
                    files_sent or [],
                    ensure_ascii=False
                ),
            )
        )

        conn.commit()

    except Exception as e:

        print(
            "LOG ERROR:",
            repr(e)
        )

    finally:

        if conn:
            conn.close()


# ============================================================
# EXCEL LOCAL STORAGE
# ============================================================

def save_excel_to_db(
    file_path,
    filename,
    session_name
):

    try:

        init_db()

        file_path = Path(
            file_path
        ).resolve()

        if not file_path.exists():

            raise FileNotFoundError(
                f"Excel file not found: {file_path}"
            )

        EXCELS_DIR.mkdir(
            parents=True,
            exist_ok=True
        )

        timestamp = utc_now().strftime(
            "%Y%m%d_%H%M%S_%f"
        )

        safe_filename = secure_filename(
            filename or "upload.xlsx"
        )

        if not safe_filename:

            safe_filename = "upload.xlsx"

        storage_filename = (
            f"{timestamp}_{safe_filename}"
        )

        storage_path = (
            EXCELS_DIR /
            storage_filename
        )

        shutil.copy2(
            str(file_path),
            str(storage_path)
        )

        conn = get_db()

        conn.execute(
            """
            INSERT INTO excel_uploads
            (
                session_name,
                filename,
                storage_path
            )
            VALUES (?, ?, ?)
            """,
            (
                session_name,
                filename or safe_filename,
                str(storage_path),
            )
        )

        conn.commit()
        conn.close()

        return True

    except Exception as e:

        print(
            "EXCEL STORAGE ERROR:",
            repr(e)
        )

        return False


# ============================================================
# CLEAN OLD EXCEL FILES
# ============================================================

def clean_old_excels():

    conn = None

    try:

        init_db()

        conn = get_db()

        cutoff = (
            utc_now()
            - datetime.timedelta(days=1)
        ).strftime(
            "%Y-%m-%d %H:%M:%S"
        )

        old = conn.execute(
            """
            SELECT *
            FROM excel_uploads
            WHERE uploaded_at < ?
            """,
            (cutoff,)
        ).fetchall()

        for ex in old:

            try:

                path = Path(
                    ex["storage_path"]
                )

                if path.exists():

                    path.unlink()

            except Exception as e:

                print(
                    "OLD EXCEL DELETE ERROR:",
                    repr(e)
                )

            conn.execute(
                """
                DELETE FROM excel_uploads
                WHERE id = ?
                """,
                (ex["id"],)
            )

        conn.commit()

    except Exception as e:

        print(
            "CLEAN EXCEL ERROR:",
            repr(e)
        )

    finally:

        if conn:
            conn.close()


# ============================================================
# SAFE FILENAME
# ============================================================

def get_safe_filename(
    filename,
    fallback
):

    if not filename:

        return fallback

    filename = secure_filename(
        Path(filename).name
    )

    if not filename:

        return fallback

    return filename


# ============================================================
# SAVE UPLOAD
# ============================================================

def save_upload(
    file_storage,
    target_dir,
    fallback_name="upload.bin"
):

    if file_storage is None:

        raise ValueError(
            "No file supplied"
        )

    filename = get_safe_filename(
        file_storage.filename,
        fallback_name
    )

    target_dir = Path(
        target_dir
    )

    target_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    path = (
        target_dir /
        filename
    )

    # حفظ أثناء request فقط
    file_storage.save(
        str(path)
    )

    if not path.exists():

        raise IOError(
            f"File was not saved: {path}"
        )

    if path.stat().st_size <= 0:

        raise IOError(
            f"Saved file is empty: {path}"
        )

    return path.resolve()


# ============================================================
# SAVE GENERATE FILES
# ============================================================

def save_generate_files(
    excel_file,
    cv_file,
    template_file,
    other_file,
    base_dir
):

    base_dir = Path(
        base_dir
    )

    base_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    excel_path = save_upload(
        excel_file,
        base_dir,
        "companies.xlsx"
    )

    cv_path = save_upload(
        cv_file,
        base_dir,
        "Lebenslauf.pdf"
    )

    template_path = save_upload(
        template_file,
        base_dir,
        "Anschreiben.docx"
    )

    other_path = save_upload(
        other_file,
        base_dir,
        "Zeugnisse.pdf"
    )

    files = [

        excel_path,
        cv_path,
        template_path,
        other_path,

    ]

    for path in files:

        if not path.exists():

            raise IOError(
                f"File missing after upload: {path}"
            )

        if path.stat().st_size <= 0:

            raise IOError(
                f"File empty after upload: {path}"
            )

    return (
        excel_path,
        cv_path,
        template_path,
        other_path,
    )


# ============================================================
# OLD VALUES
# ============================================================

def save_old_values(
    salutation,
    person_full,
    email,
    adresse_3,
    output_dir
):

    output_dir = Path(
        output_dir
    )

    output_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    content = f'''
old_salutation = {salutation!r}
old_person = {person_full!r}
old_email = {email!r}
old_adresse_3 = {adresse_3!r}
'''

    with open(
        output_dir / "saved_values.py",
        "w",
        encoding="utf-8"
    ) as f:

        f.write(content)


# ============================================================
# GENERATE LETTER
# ============================================================

def generate_letter(
    template_path,
    row,
    output_dir
):

    template_path = Path(
        template_path
    ).resolve()

    output_dir = Path(
        output_dir
    ).resolve()

    output_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    if not template_path.exists():

        raise FileNotFoundError(
            f"Template not found: {template_path}"
        )

    doc = Document(
        str(template_path)
    )

    person_full = str(
        row["person"]
    )

    firma = str(
        row["firma"]
    )

    adresse = str(
        row["adresse"]
    )

    email = str(
        row["email"]
    )

    # ------------------------------------------
    # Gender
    # ------------------------------------------

    if person_full.startswith("Herr"):

        salutation = "er "
        gender_def = True

    elif person_full.startswith("Frau"):

        salutation = "e "
        gender_def = True

    else:

        salutation = "e "
        gender_def = False

    # ------------------------------------------
    # Address
    # ------------------------------------------

    if "|" in adresse:

        adresse_1, adresse_2 = [
            x.strip()
            for x in adresse.split(
                "|",
                1
            )
        ]

    else:

        adresse_1 = adresse
        adresse_2 = adresse

    adresse_3 = (
        adresse_2[6:]
        if len(adresse_2) >= 6
        else adresse_2
    )

    # ------------------------------------------
    # Replace placeholders
    # ------------------------------------------

    for p in doc.paragraphs:

        full = "".join(
            r.text
            for r in p.runs
        )

        full = full.replace(
            "{#custom}",
            firma
        )

        full = full.replace(
            "{gender}",
            salutation
        )

        full = full.replace(
            "{zeit}",
            datetime.date.today().strftime(
                "%d.%m.%Y"
            )
        )

        if gender_def:

            full = full.replace(
                "{person}",
                person_full
            )

            full = full.replace(
                "{.}",
                person_full
            )

            full = full.replace(
                "{/custom}",
                adresse_1
            )

            full = full.replace(
                "{/custom2}",
                adresse_2
            )

            full = full.replace(
                "{adre}",
                adresse_3
            )

        else:

            full = full.replace(
                "{person}",
                "Damen und Herren"
            )

            full = full.replace(
                "{.}",
                adresse_1
            )

            full = full.replace(
                "{/custom}",
                adresse_2
            )

            full = full.replace(
                "{/custom2}",
                ""
            )

            full = full.replace(
                "{adre}",
                adresse_3
            )

        for r in p.runs:

            r.text = ""

        if len(p.runs) == 0:

            p.add_run(
                full
            )

        else:

            p.runs[0].text = full

    # ------------------------------------------
    # Safe filename
    # ------------------------------------------

    safe_name = (
        firma
        .replace("/", "_")
        .replace("\\", "_")
    )

    safe_name = get_safe_filename(
        safe_name,
        "Firma"
    )

    docx_path = (
        output_dir /
        f"{safe_name}.docx"
    )

    pdf_path = (
        output_dir /
        f"{safe_name}.pdf"
    )

    # ------------------------------------------
    # Save DOCX
    # ------------------------------------------

    doc.save(
        str(docx_path)
    )

    if not docx_path.exists():

        raise IOError(
            f"DOCX was not created: {docx_path}"
        )

    # ------------------------------------------
    # Convert DOCX -> PDF
    # ------------------------------------------

    convert(
        docx_path,
        pdf_path
    )

    # ------------------------------------------
    # Verify PDF
    # ------------------------------------------

    if not pdf_path.exists():

        raise IOError(
            f"PDF was not created: {pdf_path}"
        )

    # ------------------------------------------
    # Delete DOCX ONLY AFTER PDF SUCCESS
    # ------------------------------------------

    try:

        docx_path.unlink()

    except Exception as e:

        print(
            "DOCX DELETE WARNING:",
            repr(e)
        )

    # ------------------------------------------
    # Save values
    # ------------------------------------------

    save_old_values(
        salutation,
        person_full,
        email,
        adresse_3,
        output_dir
    )

    return pdf_path


# ============================================================
# MERGE PDF
# ============================================================

def merge_pdfs(
    cv_path,
    cover_path,
    position,
    output_path
):

    cv_path = Path(
        cv_path
    ).resolve()

    cover_path = Path(
        cover_path
    ).resolve()

    output_path = Path(
        output_path
    ).resolve()

    if not cv_path.exists():

        raise FileNotFoundError(
            f"CV not found: {cv_path}"
        )

    if not cover_path.exists():

        raise FileNotFoundError(
            f"Cover PDF not found: {cover_path}"
        )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    main_reader = PdfReader(
        str(cv_path)
    )

    insert_reader = PdfReader(
        str(cover_path)
    )

    if len(insert_reader.pages) == 0:

        raise ValueError(
            "Cover PDF contains no pages."
        )

    writer = PdfWriter()

    for i in range(
        len(main_reader.pages)
    ):

        if i == position - 1:

            writer.add_page(
                insert_reader.pages[0]
            )

        writer.add_page(
            main_reader.pages[i]
        )

    if position > len(
        main_reader.pages
    ):

        writer.add_page(
            insert_reader.pages[0]
        )

    with open(
        output_path,
        "wb"
    ) as f:

        writer.write(f)

    if not output_path.exists():

        raise IOError(
            f"Final PDF was not created: {output_path}"
        )


# ============================================================
# PATH
# ============================================================

def safe_path(p):

    if p and isinstance(
        p,
        str
    ):

        return Path(
            p
        ).resolve()

    return p
    

# ============================================================
# LOAD SAVED VALUES
# ============================================================

def load_module(
    cmp,
    saved_path
):

    saved_path = Path(
        saved_path
    ).resolve()

    if not saved_path.exists():

        raise FileNotFoundError(
            f"Saved values not found: {saved_path}"
        )

    module_name = (
        f"vals_{cmp}_{time.time_ns()}"
    )

    spec = (
        importlib.util
        .spec_from_file_location(
            module_name,
            str(saved_path)
        )
    )

    if (
        spec is None
        or spec.loader is None
    ):

        raise ImportError(
            f"Could not load {saved_path}"
        )

    module = (
        importlib.util
        .module_from_spec(spec)
    )

    sys.modules[
        module_name
    ] = module

    spec.loader.exec_module(
        module
    )

    return module


# ============================================================
# GMAIL
# ============================================================

def gmail_send(
    to_email,
    subject,
    body,
    file1_path=None,
    file2_path=None,
    file3_path=None
):

    if not GMAIL_USER:

        raise RuntimeError(
            "GMAIL_USER is not configured."
        )

    if not GMAIL_PASS:

        raise RuntimeError(
            "GMAIL_PASS is not configured."
        )

    SMTP_SERVER = "smtp.gmail.com"
    SMTP_PORT = 587

    message = MIMEMultipart()

    message["From"] = GMAIL_USER
    message["To"] = to_email
    message["Subject"] = subject

    message.attach(
        MIMEText(
            body,
            "plain"
        )
    )

    for fpath in [

        safe_path(file1_path),
        safe_path(file2_path),
        safe_path(file3_path),

    ]:

        if not fpath:
            continue

        if not fpath.exists():
            continue

        with open(
            fpath,
            "rb"
        ) as f:

            part = MIMEBase(
                "application",
                "octet-stream"
            )

            part.set_payload(
                f.read()
            )

        encoders.encode_base64(
            part
        )

        part.add_header(
            "Content-Disposition",
            f'attachment; filename="{fpath.name}"'
        )

        message.attach(
            part
        )

    server = smtplib.SMTP(
        SMTP_SERVER,
        SMTP_PORT,
        timeout=60
    )

    try:

        server.starttls()

        server.login(
            GMAIL_USER,
            GMAIL_PASS
        )

        server.sendmail(
            GMAIL_USER,
            to_email,
            message.as_string()
        )

    finally:

        server.quit()


# ============================================================
# GENERATION THREAD
# ============================================================

def generate_thread(
    excel_path,
    cv_path,
    template_path,
    other_path,
    anschreiben_pos,
    bewerbungsname,
    base_dir
):

    try:

        excel_path = Path(
            excel_path
        ).resolve()

        cv_path = Path(
            cv_path
        ).resolve()

        template_path = Path(
            template_path
        ).resolve()

        other_path = (
            Path(other_path).resolve()
            if other_path
            else None
        )

        base_dir = Path(
            base_dir
        ).resolve()

        # ------------------------------------------
        # Verify
        # ------------------------------------------

        required = [

            excel_path,
            cv_path,
            template_path,

        ]

        if other_path:
            required.append(
                other_path
            )

        for path in required:

            if not path.exists():

                raise FileNotFoundError(
                    f"Required file not found: {path}"
                )

        # ------------------------------------------
        # State
        # ------------------------------------------

        with state_lock:

            state["generating"] = True
            state["generated"] = False

            state["gen_progress"] = 0
            state["gen_log"] = []

            state["base_dir"] = str(
                base_dir
            )

            state["bewerbungsname"] = (
                bewerbungsname
            )

            state["anschreiben_pos"] = (
                anschreiben_pos
            )

            state["other"] = (
                str(other_path)
                if other_path
                else None
            )

        # ------------------------------------------
        # Excel local copy
        # ------------------------------------------

        save_excel_to_db(
            excel_path,
            excel_path.name,
            bewerbungsname
        )

        # ------------------------------------------
        # Read Excel
        # ------------------------------------------

        df = pd.read_excel(
            str(excel_path)
        )

        total = len(df)

        with state_lock:

            state["gen_total"] = total
            state["total_companies"] = total

        # ------------------------------------------
        # Empty Excel
        # ------------------------------------------

        if total == 0:

            raise ValueError(
                "Die Excel-Datei enthält keine Unternehmen."
            )

        # ------------------------------------------
        # Generate
        # ------------------------------------------

        for i, row in df.iterrows():

            cmp_num = i + 1

            out_dir = (
                base_dir /
                str(cmp_num)
            )

            out_dir.mkdir(
                parents=True,
                exist_ok=True
            )

            firma = str(
                row.get(
                    "firma",
                    ""
                )
            )

            try:

                cover_pdf = generate_letter(
                    template_path,
                    row,
                    out_dir
                )

                final_path = (
                    out_dir /
                    f"{bewerbungsname}.pdf"
                )

                merge_pdfs(
                    cv_path,
                    cover_pdf,
                    anschreiben_pos,
                    final_path
                )

                log_event(
                    bewerbungsname,
                    "generated",
                    company_num=cmp_num,
                    firma=firma,
                    status="ok"
                )

                # حذف Cover Letter المؤقت
                if cover_pdf.exists():

                    cover_pdf.unlink()

                entry = {

                    "num": cmp_num,
                    "firma": firma,
                    "status": "ok",

                }

            except Exception as e:

                print(
                    f"GENERATION ERROR #{cmp_num}:",
                    repr(e)
                )

                log_event(
                    bewerbungsname,
                    "generated",
                    company_num=cmp_num,
                    firma=firma,
                    status="error",
                    error_msg=str(e)
                )

                try:

                    with open(
                        out_dir / "skip.txt",
                        "w",
                        encoding="utf-8"
                    ) as f:

                        f.write(
                            str(e)
                        )

                except Exception:
                    pass

                entry = {

                    "num": cmp_num,
                    "firma": firma,
                    "status": "error",
                    "error": str(e),

                }

            with state_lock:

                state["gen_log"].append(
                    entry
                )

                state["gen_progress"] = (
                    cmp_num / total
                )
                # ------------------------------------------
                # Write metadata for future import
                # ------------------------------------------

                try:

                    metadata = {

                        "bewerbungsname":
                            bewerbungsname,

                        "total_companies":
                            total,

                        "other_file": (

                            other_path.name
                            if other_path
                            else None

                        ),

                        "cv_file":
                            cv_path.name,

                        "template_file":
                            template_path.name,

                        "excel_file":
                            excel_path.name,

                    }

                    with open(

                        base_dir / "metadata.json",
                        "w",
                        encoding="utf-8"

                    ) as f:

                        json.dump(
                            metadata,
                            f,
                            ensure_ascii=False,
                            indent=2
                        )

                except Exception as e:

                    print(
                        "METADATA WRITE ERROR:",
                        repr(e)
                    )

        # ------------------------------------------
        # Complete
        # ------------------------------------------

        with state_lock:
            state["generating"] = False
            state["generated"] = True
            state["gen_progress"] = 1

        # ------------------------------------------
        # Complete
        # ------------------------------------------

        with state_lock:

            state["generating"] = False
            state["generated"] = True
            state["gen_progress"] = 1

    except Exception as e:

        print(
            "FATAL GENERATION ERROR:",
            repr(e)
        )

        log_event(
            bewerbungsname,
            "generation_fatal",
            status="error",
            error_msg=str(e)
        )

        with state_lock:

            state["generating"] = False
            state["generated"] = False

            state["gen_log"].append({

                "num": 0,
                "firma": "",
                "status": "fatal",
                "error": str(e),

            })


# ============================================================
# SEND THREAD
# ============================================================

def send_thread(
    letter_path,
    delay,
    start_num,
    scheduled_dt,
    extra_path
):

    try:

        letter_path = Path(
            letter_path
        ).resolve()

        extra_path = (
            Path(extra_path).resolve()
            if extra_path
            else None
        )

        with state_lock:

            state["sending"] = True
            state["send_done"] = False
            state["interrupted_at"] = None

            state["send_log"] = []
            state["send_progress"] = 0

            state["start"] = start_num
            state["scheduled_dt"] = scheduled_dt

            state["network_error"] = False

            base_dir = Path(
                state["base_dir"]
            ).resolve()

            bewerbungsname = (
                state["bewerbungsname"]
            )

            total = (
                state["total_companies"]
            )

            state["extra"] = (
                str(extra_path)
                if extra_path
                else None
            )

        # ------------------------------------------
        # Check template
        # ------------------------------------------

        if not letter_path.exists():

            raise FileNotFoundError(
                f"Email template not found: {letter_path}"
            )

        # ------------------------------------------
        # Read template
        # ------------------------------------------

        doc = Document(
            str(letter_path)
        )

        lines = [

            p.text.strip()
            for p in doc.paragraphs
            if p.text.strip()

        ]

        if not lines:

            raise ValueError(
                "Email template is empty"
            )

        subject = lines[0]

        message_template = "\n".join(
            lines[1:]
        )

        # ------------------------------------------
        # Schedule
        # ------------------------------------------

        if scheduled_dt:

            with state_lock:

                state[
                    "waiting_scheduled"
                ] = True

            while True:

                now = datetime.datetime.now()

                remaining = (
                    scheduled_dt - now
                ).total_seconds()

                if remaining <= 0:
                    break

                time.sleep(
                    min(
                        remaining,
                        1
                    )
                )

            with state_lock:

                state[
                    "waiting_scheduled"
                ] = False

        # ------------------------------------------
        # Send
        # ------------------------------------------

        for cmp in range(
            start_num,
            total + 1
        ):

            cmp_dir = (
                base_dir /
                str(cmp)
            )

            saved_path = (
                cmp_dir /
                "saved_values.py"
            )

            pdf_path = (
                cmp_dir /
                f"{bewerbungsname}.pdf"
            )

            email = "???"

            try:

                m = load_module(
                    cmp,
                    saved_path
                )

                email = m.old_email

            except Exception:
                pass

            # ----------------------------------
            # Skip
            # ----------------------------------

            if (
                cmp_dir /
                "skip.txt"
            ).exists():

                log_event(
                    bewerbungsname,
                    "skip",
                    company_num=cmp,
                    email=email,
                    status="skip"
                )

                with state_lock:

                    state[
                        "send_log"
                    ].append({

                        "num": cmp,
                        "email": email,
                        "status": "skip",

                    })

                    state[
                        "send_progress"
                    ] = cmp / total

                continue

            # ----------------------------------
            # Send company
            # ----------------------------------

            try:

                m = load_module(
                    cmp,
                    saved_path
                )

                email = m.old_email
                gender = m.old_salutation
                person = m.old_person

                adresse_3 = (
                    m.old_adresse_3.strip()
                )

                if person == "x":

                    person = (
                        "Damen und Herren"
                    )

                letter = (

                    message_template

                    .replace(
                        "{person}",
                        person
                    )

                    .replace(
                        "{gender}",
                        gender
                    )

                    .replace(
                        "{adre}",
                        adresse_3
                    )

                    .replace(
                        "{space}",
                        "\n"
                    )

                    .replace(
                        "{2space}",
                        "\n\n"
                    )

                )

                with state_lock:

                    other_path = (
                        state.get("other")
                    )

                    current_extra = (
                        state.get("extra")
                    )

                gmail_send(
                    email,
                    subject,
                    letter,

                    str(pdf_path)
                    if pdf_path.exists()
                    else None,

                    current_extra,

                    other_path
                )

                files_sent = []

                if pdf_path.exists():

                    files_sent.append(
                        pdf_path.name
                    )

                if current_extra:

                    files_sent.append(
                        Path(
                            current_extra
                        ).name
                    )

                if other_path:

                    files_sent.append(
                        Path(
                            other_path
                        ).name
                    )

                log_event(
                    bewerbungsname,
                    "sent",
                    company_num=cmp,
                    email=email,
                    status="sent",
                    files_sent=files_sent
                )

                with state_lock:

                    state[
                        "send_log"
                    ].append({

                        "num": cmp,
                        "email": email,
                        "status": "sent",

                    })

                    state[
                        "send_progress"
                    ] = cmp / total

                time.sleep(
                    delay
                )

            except Exception as e:

                print(
                    f"SEND ERROR #{cmp}:",
                    repr(e)
                )

                if is_network_error(e):

                    log_event(
                        bewerbungsname,
                        "network_error",
                        company_num=cmp,
                        email=email,
                        status="network_error",
                        error_msg=str(e)
                    )

                    with state_lock:

                        state[
                            "send_log"
                        ].append({

                            "num": cmp,
                            "email": email,
                            "status":
                                "network_error",

                        })

                        state[
                            "interrupted_at"
                        ] = cmp

                        state[
                            "sending"
                        ] = False

                        state[
                            "network_error"
                        ] = True

                    return

                else:

                    log_event(
                        bewerbungsname,
                        "error",
                        company_num=cmp,
                        email=email,
                        status="error",
                        error_msg=str(e)
                    )

                    with state_lock:

                        state[
                            "send_log"
                        ].append({

                            "num": cmp,
                            "email": email,
                            "status": "error",

                        })

                        state[
                            "send_progress"
                        ] = cmp / total

        # ------------------------------------------
        # Complete
        # ------------------------------------------

        with state_lock:

            state["sending"] = False
            state["scheduled_dt"] = None
            state["send_done"] = True

    except Exception as e:

        print(
            "FATAL SEND ERROR:",
            repr(e)
        )

        with state_lock:

            state["sending"] = False
            state["waiting_scheduled"] = False
            state["send_done"] = False

            state[
                "send_log"
            ].append({

                "num": 0,
                "email": "",
                "status": "fatal",
                "error": str(e),

            })
            # ============================================================
# DIRECT SEND THREAD (Without Generate)
# ============================================================

def send_direct_thread():

    try:

        with state_lock:

            letter_path = Path(

                state["direct_letter_path"]

            ).resolve()

            pdf_path = (

                Path(

                    state["direct_pdf_path"]

                ).resolve()

                if state.get("direct_pdf_path")
                else None

            )

            delay = state.get("delay", 10)
            start_num = state.get("start", 1)
            scheduled_dt = state.get("scheduled_dt")

            emails_list = state.get(
                "direct_emails",
                []
            )

            total = len(emails_list)

            state["sending"] = True
            state["send_done"] = False
            state["interrupted_at"] = None
            state["send_log"] = []
            state["send_progress"] = 0
            state["network_error"] = False
            state["total_companies"] = total

        # ------------------------------------------
        # Check template
        # ------------------------------------------

        if not letter_path.exists():

            raise FileNotFoundError(

                f"Email template not found: "
                f"{letter_path}"

            )

        # ------------------------------------------
        # Read template (as-is, no replacements)
        # ------------------------------------------

        doc = Document(
            str(letter_path)
        )

        lines = [

            p.text.strip()
            for p in doc.paragraphs
            if p.text.strip()

        ]

        if not lines:

            raise ValueError(
                "Email template is empty"
            )

        subject = lines[0]

        message_body = "\n".join(
            lines[1:]
        )

        # ------------------------------------------
        # Schedule wait
        # ------------------------------------------

        if scheduled_dt:

            with state_lock:

                state[
                    "waiting_scheduled"
                ] = True

            while True:

                now = datetime.datetime.now()

                remaining = (

                    scheduled_dt - now

                ).total_seconds()

                if remaining <= 0:
                    break

                time.sleep(
                    min(remaining, 1)
                )

            with state_lock:

                state[
                    "waiting_scheduled"
                ] = False

        # ------------------------------------------
        # Send emails
        # ------------------------------------------

        for i in range(

            start_num - 1,
            total

        ):

            email = emails_list[i]
            cmp_num = i + 1

            try:

                gmail_send(

                    email,
                    subject,
                    message_body,

                    str(pdf_path)
                    if pdf_path
                    and pdf_path.exists()
                    else None,

                    None,
                    None

                )

                log_event(

                    "Direct_Send",
                    "sent",

                    company_num=cmp_num,
                    email=email,
                    status="sent",

                    files_sent=(
                        [pdf_path.name]
                        if pdf_path
                        else []
                    )

                )

                with state_lock:

                    state[
                        "send_log"
                    ].append({

                        "num": cmp_num,
                        "email": email,
                        "status": "sent",

                    })

                    state[
                        "send_progress"
                    ] = cmp_num / total

                time.sleep(delay)

            except Exception as e:

                print(

                    f"DIRECT SEND ERROR "
                    f"#{cmp_num}:",
                    repr(e)

                )

                if is_network_error(e):

                    log_event(

                        "Direct_Send",
                        "network_error",

                        company_num=cmp_num,
                        email=email,
                        status="network_error",
                        error_msg=str(e)

                    )

                    with state_lock:

                        state[
                            "send_log"
                        ].append({

                            "num": cmp_num,
                            "email": email,
                            "status":
                                "network_error",

                        })

                        state[
                            "interrupted_at"
                        ] = cmp_num

                        state[
                            "sending"
                        ] = False

                        state[
                            "network_error"
                        ] = True

                    return

                else:

                    log_event(

                        "Direct_Send",
                        "error",

                        company_num=cmp_num,
                        email=email,
                        status="error",
                        error_msg=str(e)

                    )

                    with state_lock:

                        state[
                            "send_log"
                        ].append({

                            "num": cmp_num,
                            "email": email,
                            "status": "error",

                        })

                        state[
                            "send_progress"
                        ] = cmp_num / total

        # ------------------------------------------
        # Complete
        # ------------------------------------------

        with state_lock:

            state["sending"] = False
            state["scheduled_dt"] = None
            state["send_done"] = True

    except Exception as e:

        print(

            "FATAL DIRECT SEND ERROR:",
            repr(e)

        )

        with state_lock:

            state["sending"] = False
            state["waiting_scheduled"] = False
            state["send_done"] = False

            state[
                "send_log"
            ].append({

                "num": 0,
                "email": "",
                "status": "fatal",
                "error": str(e),

            })


# ============================================================
# INDEX
# ============================================================

@app.route("/")
def index():

    return render_template(
        "index.html"
    )


# ============================================================
# LOGIN
# ============================================================

@app.route(
    "/api/login",
    methods=["POST"]
)
def login():

    data = request.json or {}

    code = data.get(
        "code",
        ""
    )

    if code == ADMIN_CODE:

        with state_lock:

            state["logged_in"] = True
            state["is_admin"] = True

        return jsonify({

            "success": True,
            "is_admin": True,

        })

    if code == ACCESS_CODE:

        with state_lock:

            state["logged_in"] = True
            state["is_admin"] = False

        return jsonify({

            "success": True,
            "is_admin": False,

        })

    return jsonify({

        "success": False,
        "error": "Falscher Code",

    }), 401


# ============================================================
# STATE
# ============================================================

@app.route(
    "/api/state",
    methods=["GET"]
)
def get_state():

    with state_lock:

        scheduled = state.get(
            "scheduled_dt"
        )

        return jsonify({

            "logged_in":
                state["logged_in"],

            "is_admin":
                state["is_admin"],

            "generating":
                state["generating"],

            "generated":
                state["generated"],

            "sending":
                state["sending"],

            "send_done":
                state["send_done"],

            "interrupted_at":
                state["interrupted_at"],

            "bewerbungsname":
                state["bewerbungsname"],

            "total_companies":
                state["total_companies"],

            "waiting_scheduled":
                state[
                    "waiting_scheduled"
                ],

            "scheduled_dt":
                (
                    scheduled.isoformat()
                    if scheduled
                    else None
                ),

        })


# ============================================================
# GENERATE
# ============================================================

@app.route(
    "/api/generate",
    methods=["POST"]
)
def generate():

    with state_lock:

        if not state["logged_in"]:

            return jsonify({

                "success": False,
                "error": "Not logged in",

            }), 403

        if state["generating"]:

            return jsonify({

                "success": False,
                "error":
                    "Generierung läuft bereits",

            }), 400

    # ------------------------------------------
    # Receive files
    # ------------------------------------------

    excel_file = request.files.get(
        "excel"
    )

    cv_file = request.files.get(
        "cv"
    )

    template_file = request.files.get(
        "template"
    )

    other_file = request.files.get(
        "other"
    )

    # ------------------------------------------
    # Validate
    # ------------------------------------------

    if not excel_file:

        return jsonify({

            "success": False,
            "error": "Excel fehlt",

        }), 400

    if not cv_file:

        return jsonify({

            "success": False,
            "error": "Lebenslauf fehlt",

        }), 400

    if not template_file:

        return jsonify({

            "success": False,
            "error": "Anschreiben fehlt",

        }), 400

    if not other_file:

        return jsonify({

            "success": False,
            "error":
                "Zeugnisse und Zertifikate fehlen",

        }), 400

    # ------------------------------------------
    # Position
    # ------------------------------------------

    try:

        anschreiben_pos = int(
            request.form.get(
                "position",
                2
            )
        )

        if anschreiben_pos < 1:

            anschreiben_pos = 1

    except Exception:

        anschreiben_pos = 2

    # ------------------------------------------
    # Application name
    # ------------------------------------------

    bewerbungsname = (
        request.form.get(
            "bewerbungsname",
            "Bewerbung"
        ).strip()
    )

    if not bewerbungsname:

        bewerbungsname = "Bewerbung"

    safe_bewerbungsname = secure_filename(
        bewerbungsname
    )

    if not safe_bewerbungsname:

        safe_bewerbungsname = "Bewerbung"

    # ------------------------------------------
    # Directory
    # ------------------------------------------

    base_dir = (
        APPLICATIONS_DIR /
        safe_bewerbungsname
    ).resolve()

    try:

        if base_dir.exists():

            shutil.rmtree(
                str(base_dir)
            )

        base_dir.mkdir(
            parents=True,
            exist_ok=True
        )

        # مهم:
        # حفظ الملفات هنا داخل request
        # وليس داخل background thread

        (
            excel_path,
            cv_path,
            template_path,
            other_path
        ) = save_generate_files(

            excel_file,
            cv_file,
            template_file,
            other_file,

            base_dir
        )

    except Exception as e:

        print(
            "UPLOAD ERROR:",
            repr(e)
        )

        return jsonify({

            "success": False,

            "error":
                "Fehler beim Speichern der Dateien: "
                + str(e),

        }), 500

    # ------------------------------------------
    # State
    # ------------------------------------------

    with state_lock:

        state["generating"] = True
        state["generated"] = False

        state["gen_progress"] = 0
        state["gen_log"] = []
        state["gen_total"] = 0

        state["base_dir"] = str(
            base_dir
        )

        state["bewerbungsname"] = (
            safe_bewerbungsname
        )

        state["anschreiben_pos"] = (
            anschreiben_pos
        )

        state["other"] = str(
            other_path
        )

        state["extra"] = None

    # ------------------------------------------
    # Save Excel copy
    # ------------------------------------------

    save_excel_to_db(

        excel_path,

        excel_file.filename,

        safe_bewerbungsname

    )

    # ------------------------------------------
    # Thread
    # ------------------------------------------

    thread = threading.Thread(

        target=generate_thread,

        args=(

            str(excel_path),
            str(cv_path),
            str(template_path),
            str(other_path),

            anschreiben_pos,

            safe_bewerbungsname,

            str(base_dir),

        ),

        daemon=True
    )

    thread.start()

    return jsonify({

        "success": True

    }), 200


# ============================================================
# GENERATE STATUS
# ============================================================

@app.route(
    "/api/generate/status",
    methods=["GET"]
)
def generate_status():

    with state_lock:

        return jsonify({

            "generating":
                state["generating"],

            "generated":
                state["generated"],

            "progress":
                state["gen_progress"],

            "log":
                state["gen_log"],

            "total":
                state["gen_total"],

        })
        
# ============================================================
# IMPORT GENERATED FILES (ZIP)
# ============================================================

@app.route(
    "/api/generate/import",
    methods=["POST"]
)
def import_generated():

    with state_lock:

        if not state["logged_in"]:

            return jsonify({

                "success": False,
                "error": "Not logged in",

            }), 403

        if state["generating"] or state["sending"]:

            return jsonify({

                "success": False,
                "error": "Es läuft bereits ein Prozess",

            }), 400

    # ------------------------------------------
    # Receive ZIP
    # ------------------------------------------

    zip_file = request.files.get(
        "zip"
    )

    if not zip_file:

        return jsonify({

            "success": False,
            "error": "ZIP Datei fehlt",

        }), 400

    # ------------------------------------------
    # Temp extraction dir
    # ------------------------------------------

    temp_dir = (

        APPLICATIONS_DIR
        / f"temp_import_{int(time.time())}"

    )

    temp_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    safe_bewerbungsname = "Bewerbung"
    total_companies = 0

    try:

        # ----------------------------------
        # Save ZIP
        # ----------------------------------

        temp_zip_path = save_upload(

            zip_file,
            temp_dir,
            "import.zip"

        )

        # ----------------------------------
        # Extract
        # ----------------------------------

        with zipfile.ZipFile(

            str(temp_zip_path),
            "r"

        ) as zipf:

            zipf.extractall(
                str(temp_dir)
            )

        # ----------------------------------
        # Find root folder inside ZIP
        # ----------------------------------

        root_folder = None

        for item in temp_dir.iterdir():

            if item.is_dir() and item.name != "__MACOSX":

                root_folder = item
                break

        if not root_folder:

            # Files directly in temp_dir
            root_folder = temp_dir

        bewerbungsname = root_folder.name

        safe_bewerbungsname = secure_filename(
            bewerbungsname
        )

        if not safe_bewerbungsname:

            safe_bewerbungsname = "Bewerbung"

        # ----------------------------------
        # Target dir
        # ----------------------------------

        target_dir = (

            APPLICATIONS_DIR
            / safe_bewerbungsname

        ).resolve()

        if target_dir.exists():

            shutil.rmtree(
                str(target_dir)
            )

        # ----------------------------------
        # Move extracted to target
        # ----------------------------------

        if root_folder != temp_dir:

            shutil.move(

                str(root_folder),
                str(target_dir)

            )

        else:

            target_dir.mkdir(
                parents=True,
                exist_ok=True
            )

            for item in temp_dir.iterdir():

                if item.name == "import.zip":
                    continue

                shutil.move(

                    str(item),
                    str(target_dir / item.name)

                )

        # ----------------------------------
        # Read metadata if exists
        # ----------------------------------

        metadata_path = (
            target_dir / "metadata.json"
        )

        metadata = None

        if metadata_path.exists():

            try:

                with open(

                    metadata_path,
                    "r",
                    encoding="utf-8"

                ) as f:

                    metadata = json.load(f)

            except Exception:
                pass

        # ----------------------------------
        # Count company folders
        # ----------------------------------

        company_folders = [

            d for d in target_dir.iterdir()

            if d.is_dir() and d.name.isdigit()

        ]

        total_companies = len(
            company_folders
        )

        if total_companies == 0:

            return jsonify({

                "success": False,
                "error":
                    "Keine Unternehmen im ZIP gefunden",

            }), 400

        # ----------------------------------
        # Find other file (Zeugnisse)
        # ----------------------------------

        other_path = None

        if metadata and metadata.get("other_file"):

            other_candidate = (

                target_dir
                / metadata["other_file"]

            )

            if other_candidate.exists():

                other_path = str(
                    other_candidate.resolve()
                )

        if not other_path:

            pdf_files = list(
                target_dir.glob("*.pdf")
            )

            # Exclude CV
            cv_name = None

            if metadata:
                cv_name = metadata.get("cv_file")

            if cv_name:

                pdf_files = [

                    f for f in pdf_files
                    if f.name != cv_name

                ]

            else:

                pdf_files = [

                    f for f in pdf_files
                    if "lebenslauf" not in f.name.lower()

                ]

            if pdf_files:

                other_path = str(
                    pdf_files[0].resolve()
                )

        # ----------------------------------
        # Update state
        # ----------------------------------

        with state_lock:

            state["generated"] = True
            state["generating"] = False

            state["base_dir"] = str(
                target_dir
            )

            state["bewerbungsname"] = (
                safe_bewerbungsname
            )

            state["total_companies"] = (
                total_companies
            )

            state["other"] = other_path
            state["extra"] = None

            state["anschreiben_pos"] = 2

            state["send_done"] = False
            state["interrupted_at"] = None

            state["send_progress"] = 0
            state["send_log"] = []

            state["gen_progress"] = 1
            state["gen_log"] = []
            state["gen_total"] = total_companies

            state["scheduled_dt"] = None
            state["waiting_scheduled"] = False
            state["network_error"] = False

    except Exception as e:

        print(
            "IMPORT ERROR:",
            repr(e)
        )

        return jsonify({

            "success": False,

            "error":
                "Fehler beim Import: "
                + str(e),

        }), 500

    finally:

        try:

            shutil.rmtree(
                str(temp_dir),
                ignore_errors=True
            )

        except Exception:
            pass

    return jsonify({

        "success": True,

        "bewerbungsname":
            safe_bewerbungsname,

        "total_companies":
            total_companies,

    })

# ============================================================
# EXPORT GENERATED FILES
# ============================================================

@app.route(
    "/api/generate/export",
    methods=["GET"]
)
def export_generated():

    with state_lock:

        if not state["logged_in"]:

            return jsonify({

                "success": False,
                "error": "Not logged in",

            }), 403

        if not state["generated"]:

            return jsonify({

                "success": False,
                "error": "Keine generierten Dateien vorhanden",

            }), 400

        base_dir = Path(
            state["base_dir"]
        ).resolve()

        bewerbungsname = state[
            "bewerbungsname"
        ]

    # ------------------------------------------
    # Check folder exists
    # ------------------------------------------

    if not base_dir.exists():

        return jsonify({

            "success": False,
            "error": "Ordner nicht gefunden",

        }), 404

    # ------------------------------------------
    # Create ZIP in temp file
    # ------------------------------------------

    temp_fd, temp_path = tempfile.mkstemp(
        suffix=".zip"
    )

    os.close(
        temp_fd
    )

    try:

        with zipfile.ZipFile(

            temp_path,
            "w",
            zipfile.ZIP_DEFLATED

        ) as zipf:

            for root, dirs, files in os.walk(
                str(base_dir)
            ):

                for file in files:

                    file_path = (
                        Path(root) / file
                    )

                    # Keep folder structure:
                    # bewerbungsname/1/file.pdf
                    arcname = (

                        Path(bewerbungsname)
                        / file_path.relative_to(base_dir)

                    )

                    zipf.write(
                        str(file_path),
                        str(arcname)
                    )

    except Exception as e:

        try:
            os.unlink(temp_path)
        except Exception:
            pass

        print(
            "EXPORT ZIP ERROR:",
            repr(e)
        )

        return jsonify({

            "success": False,
            "error":
                "Fehler beim Erstellen der ZIP-Datei: "
                + str(e),

        }), 500

    # ------------------------------------------
    # Cleanup after sending
    # ------------------------------------------

    @after_this_request
    def cleanup(response):

        try:
            os.unlink(temp_path)
        except Exception:
            pass

        return response

    return send_file(

        temp_path,

        as_attachment=True,

        download_name=(
            f"{bewerbungsname}.zip"
        )

    )
# ============================================================
# SEND
# ============================================================

@app.route(
    "/api/send",
    methods=["POST"]
)
def send():

    with state_lock:

        if not state["logged_in"]:

            return jsonify({

                "success": False,
                "error": "Not logged in",

            }), 403

        if not state["generated"]:

            return jsonify({

                "success": False,
                "error":
                    "Bitte zuerst Anschreiben generieren",

            }), 400

        if state["sending"]:

            return jsonify({

                "success": False,
                "error":
                    "Senden läuft bereits",

            }), 400

    letter_file = request.files.get(
        "letter"
    )

    extra_file = request.files.get(
        "extra"
    )

    if not letter_file:

        return jsonify({

            "success": False,
            "error":
                "Email Template erforderlich",

        }), 400

    # ------------------------------------------
    # Delay
    # ------------------------------------------

    try:

        delay = int(
            request.form.get(
                "delay",
                10
            )
        )

    except Exception:

        delay = 10

    if delay < 0:
        delay = 0

    # ------------------------------------------
    # Start
    # ------------------------------------------

    try:

        start_num = int(
            request.form.get(
                "start",
                1
            )
        )

    except Exception:

        start_num = 1

    if start_num < 1:
        start_num = 1

    # ------------------------------------------
    # Schedule
    # ------------------------------------------

    schedule_str = request.form.get(
        "scheduled_dt",
        ""
    )

    scheduled_dt = None

    if schedule_str:

        try:

            scheduled_dt = (
                datetime.datetime
                .fromisoformat(
                    schedule_str
                )
            )

            if (
                scheduled_dt
                <= datetime.datetime.now()
            ):

                return jsonify({

                    "success": False,
                    "error":
                        "Die gewählte Zeit liegt "
                        "in der Vergangenheit",

                }), 400

        except Exception:

            return jsonify({

                "success": False,
                "error":
                    "Ungültiges Datum/Zeit Format",

            }), 400

    # ------------------------------------------
    # Base dir
    # ------------------------------------------

    with state_lock:

        if not state["base_dir"]:

            return jsonify({

                "success": False,
                "error":
                    "Application folder not found",

            }), 400

        base_dir = Path(
            state["base_dir"]
        ).resolve()

    # ------------------------------------------
    # Save email template
    # ------------------------------------------

    try:

        letter_path = save_upload(

            letter_file,

            base_dir,

            "Email_Template.docx"

        )

        if (
            extra_file
            and extra_file.filename
        ):

            extra_path = save_upload(

                extra_file,

                base_dir,

                "Anhang.pdf"

            )

        else:

            extra_path = None

    except Exception as e:

        print(
            "SEND UPLOAD ERROR:",
            repr(e)
        )

        return jsonify({

            "success": False,

            "error":
                "Fehler beim Speichern "
                "der Email-Dateien: "
                + str(e),

        }), 500

    # ------------------------------------------
    # State
    # ------------------------------------------

    with state_lock:

        state["sending"] = True
        state["send_done"] = False

        state["interrupted_at"] = None

        state["send_progress"] = 0
        state["send_log"] = []

        state["delay"] = delay
        state["start"] = start_num

        state["scheduled_dt"] = scheduled_dt

        state["network_error"] = False

        state["extra"] = (
            str(extra_path)
            if extra_path
            else None
        )

    # ------------------------------------------
    # Thread
    # ------------------------------------------

    thread = threading.Thread(

        target=send_thread,

        args=(

            str(letter_path),

            delay,

            start_num,

            scheduled_dt,

            str(extra_path)
            if extra_path
            else None,

        ),

        daemon=True
    )

    thread.start()

    return jsonify({

        "success": True

    }), 200
    
# ============================================================
# DIRECT SEND (Without Generate)
# ============================================================

@app.route(
    "/api/send/direct",
    methods=["POST"]
)
def send_direct():

    with state_lock:

        if not state["logged_in"]:

            return jsonify({

                "success": False,
                "error": "Not logged in",

            }), 403

        if state["sending"]:

            return jsonify({

                "success": False,
                "error":
                    "Senden läuft bereits",

            }), 400

    # ------------------------------------------
    # Receive files
    # ------------------------------------------

    excel_file = request.files.get(
        "excel"
    )

    letter_file = request.files.get(
        "letter"
    )

    pdf_file = request.files.get(
        "pdf"
    )

    if not excel_file:

        return jsonify({

            "success": False,
            "error":
                "Excel (Emails) fehlt",

        }), 400

    if not letter_file:

        return jsonify({

            "success": False,
            "error":
                "Email Template (DOCX) fehlt",

        }), 400

    if not pdf_file:

        return jsonify({

            "success": False,
            "error":
                "PDF Anhang fehlt",

        }), 400

    # ------------------------------------------
    # Delay
    # ------------------------------------------

    try:

        delay = int(
            request.form.get(
                "delay",
                10
            )
        )

    except Exception:

        delay = 10

    if delay < 0:
        delay = 0

    # ------------------------------------------
    # Start
    # ------------------------------------------

    try:

        start_num = int(
            request.form.get(
                "start",
                1
            )
        )

    except Exception:

        start_num = 1

    if start_num < 1:
        start_num = 1

    # ------------------------------------------
    # Schedule
    # ------------------------------------------

    schedule_str = request.form.get(
        "scheduled_dt",
        ""
    )

    scheduled_dt = None

    if schedule_str:

        try:

            scheduled_dt = (

                datetime.datetime
                .fromisoformat(
                    schedule_str
                )

            )

            if (

                scheduled_dt
                <= datetime.datetime.now()

            ):

                return jsonify({

                    "success": False,
                    "error":
                        "Die gewählte Zeit "
                        "liegt in der Vergangenheit",

                }), 400

        except Exception:

            return jsonify({

                "success": False,
                "error":
                    "Ungültiges Datum/Zeit Format",

            }), 400

    # ------------------------------------------
    # Save files to temp dir
    # ------------------------------------------

    temp_dir = (

        APPLICATIONS_DIR
        / f"direct_send_{int(time.time())}"

    )

    temp_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    try:

        excel_path = save_upload(

            excel_file,
            temp_dir,
            "emails.xlsx"

        )

        letter_path = save_upload(

            letter_file,
            temp_dir,
            "Email_Template.docx"

        )

        pdf_path = save_upload(

            pdf_file,
            temp_dir,
            "Anhang.pdf"

        )

    except Exception as e:

        print(
            "DIRECT SEND UPLOAD ERROR:",
            repr(e)
        )

        shutil.rmtree(
            str(temp_dir),
            ignore_errors=True
        )

        return jsonify({

            "success": False,

            "error":
                "Fehler beim Speichern: "
                + str(e),

        }), 500

    # ------------------------------------------
    # Read emails from Excel
    # ------------------------------------------

    try:

        df = pd.read_excel(
            str(excel_path)
        )

        # Find email column
        email_col = None

        for col in df.columns:

            col_lower = str(
                col
            ).lower()

            if col_lower in (

                "email",
                "e-mail",
                "mail",
                "e_mail",
                "emails",
                "e-mail-adresse",
                "email adresse",

            ):

                email_col = col
                break

        if not email_col:

            # Use first column
            email_col = df.columns[0]

        emails_list = [

            str(e).strip()
            for e in df[email_col].tolist()
            if str(e).strip()
            and str(e).strip().lower() != "nan"

        ]

        if not emails_list:

            shutil.rmtree(
                str(temp_dir),
                ignore_errors=True
            )

            return jsonify({

                "success": False,
                "error":
                    "Keine Emails in der "
                    "Excel-Datei gefunden",

            }), 400

        if start_num > len(emails_list):

            shutil.rmtree(
                str(temp_dir),
                ignore_errors=True
            )

            return jsonify({

                "success": False,
                "error":
                    "Startnummer ist größer "
                    "als die Anzahl der Emails",

            }), 400

    except Exception as e:

        print(
            "EXCEL READ ERROR:",
            repr(e)
        )

        shutil.rmtree(
            str(temp_dir),
            ignore_errors=True
        )

        return jsonify({

            "success": False,

            "error":
                "Fehler beim Lesen der Excel: "
                + str(e),

        }), 500

    # ------------------------------------------
    # State
    # ------------------------------------------

    with state_lock:

        state["direct_mode"] = True
        state["direct_emails"] = emails_list
        state["direct_letter_path"] = str(letter_path)
        state["direct_pdf_path"] = str(pdf_path)

        state["sending"] = True
        state["send_done"] = False
        state["interrupted_at"] = None

        state["send_progress"] = 0
        state["send_log"] = []

        state["delay"] = delay
        state["start"] = start_num
        state["scheduled_dt"] = scheduled_dt

        state["network_error"] = False
        state["waiting_scheduled"] = False

        state["total_companies"] = len(
            emails_list
        )

        state["base_dir"] = str(temp_dir)
        state["bewerbungsname"] = "Direct_Send"

    # ------------------------------------------
    # Thread
    # ------------------------------------------

    thread = threading.Thread(

        target=send_direct_thread,

        daemon=True

    )

    thread.start()

    return jsonify({

        "success": True

    }), 200


# ============================================================
# SEND STATUS
# ============================================================

@app.route(
    "/api/send/status",
    methods=["GET"]
)
def send_status():

    with state_lock:

        scheduled = state.get(
            "scheduled_dt"
        )

        return jsonify({

            "sending":
                state["sending"],

            "send_done":
                state["send_done"],

            "interrupted_at":
                state["interrupted_at"],

            "progress":
                state["send_progress"],

            "log":
                state["send_log"],

            "total":
                state["total_companies"],

            "waiting_scheduled":
                state["waiting_scheduled"],

            "scheduled_dt":
                (
                    scheduled.isoformat()
                    if scheduled
                    else None
                ),

            "bewerbungsname":
                state["bewerbungsname"],

            "direct_mode":
                state.get(
                    "direct_mode",
                    False
                ),

        })


# ============================================================
# RESUME SEND
# ============================================================

@app.route(
    "/api/send/resume",
    methods=["POST"]
)
def resume_send():

    with state_lock:

        if not state["logged_in"]:

            return jsonify({

                "success": False,
                "error": "Not logged in",

            }), 403

        is_direct = state.get(
            "direct_mode",
            False
        )

        if is_direct:

            total = len(

                state.get(
                    "direct_emails",
                    []
                )

            )

        else:

            if not state["base_dir"]:

                return jsonify({

                    "success": False,
                    "error":
                        "Application folder not found",

                }), 400

            base_dir = Path(
                state["base_dir"]
            ).resolve()

            total = state[
                "total_companies"
            ]

            delay = state["delay"]
            extra_path = state.get("extra")

    data = request.json or {}

    try:

        resume_from = int(
            data.get(
                "resume_from",
                1
            )
        )

    except Exception:

        resume_from = 1

    if resume_from < 1:
        resume_from = 1

    if total and resume_from > total:

        return jsonify({

            "success": False,
            "error":
                "Ungültige Startnummer",

        }), 400

    # ------------------------------------------
    # Direct mode resume
    # ------------------------------------------

    if is_direct:

        with state_lock:

            state["interrupted_at"] = None
            state["network_error"] = False

            state["sending"] = True
            state["send_done"] = False

            state["start"] = resume_from

            state["send_progress"] = 0
            state["send_log"] = []

            state["scheduled_dt"] = None
            state["waiting_scheduled"] = False

        thread = threading.Thread(

            target=send_direct_thread,
            daemon=True

        )

        thread.start()

        return jsonify({
            "success": True
        })

    # ------------------------------------------
    # Normal mode resume (existing logic)
    # ------------------------------------------

    letter_files = list(
        base_dir.glob(
            "Email_Template.docx"
        )
    )

    if not letter_files:

        letter_files = list(
            base_dir.glob("*.docx")
        )

    if not letter_files:

        return jsonify({

            "success": False,
            "error":
                "Email Template nicht gefunden",

        }), 400

    letter_path = (
        letter_files[0].resolve()
    )

    with state_lock:

        state["interrupted_at"] = None
        state["network_error"] = False

        state["sending"] = True
        state["send_done"] = False

        state["start"] = resume_from

        state["send_progress"] = 0
        state["send_log"] = []

    thread = threading.Thread(

        target=send_thread,

        args=(

            str(letter_path),
            delay,
            resume_from,
            None,
            extra_path,

        ),

        daemon=True

    )

    thread.start()

    return jsonify({
        "success": True
    })


# ============================================================
# RESET
# ============================================================

@app.route(
    "/api/reset",
    methods=["POST"]
)
def reset():

    reset_state()

    return jsonify({

        "success": True

    })


# ============================================================
# DASHBOARD
# ============================================================

@app.route(
    "/api/dashboard",
    methods=["GET"]
)
def dashboard():

    with state_lock:

        if (
            not state["logged_in"]
            or not state["is_admin"]
        ):

            return jsonify({

                "error":
                    "Access denied"

            }), 403

    session_filter = request.args.get(
        "session",
        "Alle"
    )

    hours_filter = request.args.get(
        "hours",
        type=int
    )

    init_db()

    conn = get_db()

    query = """
        SELECT *
        FROM bewerber_logs
    """

    conditions = []
    params = []

    if hours_filter:

        cutoff = (

            utc_now()
            - datetime.timedelta(
                hours=hours_filter
            )

        ).strftime(
            "%Y-%m-%d %H:%M:%S"
        )

        conditions.append(
            "created_at >= ?"
        )

        params.append(
            cutoff
        )

    if conditions:

        query += (
            " WHERE "
            + " AND ".join(
                conditions
            )
        )

    query += """
        ORDER BY created_at DESC
        LIMIT 1000
    """

    rows = conn.execute(
        query,
        params
    ).fetchall()

    conn.close()

    # ------------------------------------------
    # Stats
    # ------------------------------------------

    sent_n = sum(

        1
        for r in rows
        if r["event_type"] == "sent"

    )

    skip_n = sum(

        1
        for r in rows
        if r["event_type"] == "skip"

    )

    error_n = sum(

        1
        for r in rows
        if r["event_type"]
        in (
            "error",
            "network_error",
            "generation_fatal",
        )

    )

    gen_ok = sum(

        1
        for r in rows

        if (

            r["event_type"]
            == "generated"

            and

            r["status"]
            == "ok"

        )

    )

    # ------------------------------------------
    # Sessions
    # ------------------------------------------

    sessions = list(
        set(

            r["session_name"]

            for r in rows

            if r["session_name"]

        )
    )

    # ------------------------------------------
    # Companies
    # ------------------------------------------

    companies = defaultdict(

        lambda: {

            "session": "",
            "generiert": "—",
            "gesendet": "Nein",
            "email_firma": "",
            "fehler": "—",
            "zeit": "",

        }
    )

    for r in sorted(

        rows,

        key=lambda x:
            x["created_at"]

    ):

        key = (

            r["session_name"],
            r["company_num"],

        )

        e = companies[key]

        e["session"] = (
            r["session_name"]
            or ""
        )

        if r["event_type"] == "generated":

            e["generiert"] = (

                "Ja"

                if r["status"] == "ok"

                else "Nein"

            )

            e["email_firma"] = (
                r["firma"]
                or ""
            )

            e["zeit"] = (

                r["created_at"][:16]

                if r["created_at"]

                else ""

            )

        if r["event_type"] == "sent":

            e["gesendet"] = "Ja"

            e["email_firma"] = (

                r["email"]
                or e["email_firma"]

            )

            e["zeit"] = (

                r["created_at"][:16]

                if r["created_at"]

                else ""

            )

        if r["event_type"] == "skip":

            e["gesendet"] = "Nein"

        if r["event_type"] in (

            "error",
            "network_error",
            "generation_fatal",

        ):

            e["fehler"] = "Ja"

    # ------------------------------------------
    # Filter
    # ------------------------------------------

    if session_filter != "Alle":

        companies = {

            k: v

            for k, v
            in companies.items()

            if k[0] == session_filter

        }

    # ------------------------------------------
    # Company list
    # ------------------------------------------

    company_list = []

    for (

        session,
        cmp_num

    ), e in sorted(

        companies.items(),

        key=lambda x: (
            x[0][1]
            or 0
        )

    ):

        company_list.append({

            "session":
                e["session"],

            "cmp_num":
                cmp_num,

            "generiert":
                e["generiert"],

            "gesendet":
                e["gesendet"],

            "email_firma":
                e["email_firma"],

            "fehler":
                e["fehler"],

            "zeit":
                e["zeit"],

        })

    return jsonify({

        "stats": {

            "generated":
                gen_ok,

            "sent":
                sent_n,

            "skipped":
                skip_n,

            "errors":
                error_n,

        },

        "sessions":
            sessions,

        "companies":
            company_list,

    })


# ============================================================
# DASHBOARD EXCELS
# ============================================================

@app.route(
    "/api/dashboard/excels",
    methods=["GET"]
)
def dashboard_excels():

    with state_lock:

        if (
            not state["logged_in"]
            or not state["is_admin"]
        ):

            return jsonify({

                "error":
                    "Access denied"

            }), 403

    clean_old_excels()

    init_db()

    conn = get_db()

    rows = conn.execute(
        """
        SELECT *
        FROM excel_uploads
        ORDER BY uploaded_at DESC
        LIMIT 200
        """
    ).fetchall()

    conn.close()

    session_filter = request.args.get(
        "session",
        "Alle"
    )

    files = []

    for r in rows:

        if (

            session_filter != "Alle"

            and

            r["session_name"]
            != session_filter

        ):

            continue

        files.append({

            "id":
                r["id"],

            "filename":
                r["filename"],

            "session_name":
                r["session_name"],

            "uploaded_at": (

                r["uploaded_at"][:16]

                if r["uploaded_at"]

                else ""

            ),

        })

    all_sessions = list(
        set(

            r["session_name"]

            for r in rows

            if r["session_name"]

        )
    )

    return jsonify({

        "files":
            files,

        "sessions":
            all_sessions,

    })


# ============================================================
# DOWNLOAD EXCEL
# ============================================================

@app.route(
    "/api/dashboard/excel/download/<int:file_id>"
)
def download_excel(file_id):

    with state_lock:

        if (
            not state["logged_in"]
            or not state["is_admin"]
        ):

            return jsonify({

                "error":
                    "Access denied"

            }), 403

    init_db()

    conn = get_db()

    row = conn.execute(
        """
        SELECT *
        FROM excel_uploads
        WHERE id = ?
        """,
        (file_id,)
    ).fetchone()

    conn.close()

    if not row:

        return jsonify({

            "error":
                "File not found"

        }), 404

    path = Path(
        row["storage_path"]
    ).resolve()

    if not path.exists():

        return jsonify({

            "error":
                "File no longer exists"

        }), 404

    return send_file(

        str(path),

        as_attachment=True,

        download_name=
            row["filename"]

    )


# ============================================================
# ERROR HANDLER
# ============================================================

@app.errorhandler(413)
def request_entity_too_large(error):

    return jsonify({

        "success": False,

        "error":
            "File upload is too large."

    }), 413


# ============================================================
# RUN
# ============================================================

if __name__ == "__main__":

    print()
    print("=" * 60)
    print("Bewerber Flask")
    print("=" * 60)
    print(
        "BASE_DIR:",
        BASE_DIR
    )
    print(
        "DATA_DIR:",
        DATA_DIR
    )
    print(
        "APPLICATIONS_DIR:",
        APPLICATIONS_DIR
    )
    print(
        "EXCELS_DIR:",
        EXCELS_DIR
    )
    print(
        "DATABASE:",
        LOGS_DB
    )

    try:

        print(
            "LibreOffice:",
            get_libreoffice()
        )

    except Exception as e:

        print(
            "LibreOffice WARNING:",
            e
        )

    print("=" * 60)
    print()

    app.run(
        host="0.0.0.0",
        port=5000,
        debug=False,
        threaded=True
    )